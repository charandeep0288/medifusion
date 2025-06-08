# Medical Document Parser with OpenAI GPT and AWS Textract
# For Google Colab - Handles Images, PDFs, and Word Documents

# # Install required packages
# !pip install openai
# !pip install python-docx
# !pip install PyPDF2
# !pip install pill
# !pip install pdf2image
# !pip install boto3
# !apt-get update
# !apt-get install poppler-utils

import os
import json
import base64
import io
import mimetypes
from typing import Dict, Any, List
import os
import json
import base64
import io
import mimetypes
from typing import Dict, Any, List
from pathlib import Path
from PIL import Image
from pdf2image import convert_from_path, convert_from_bytes
import PyPDF2
from docx import Document
from openai import OpenAI
import requests
import boto3
from botocore.exceptions import ClientError

class MedicalDocumentParser:
    def __init__(self, openai_api_key: str, aws_access_key: str, aws_secret_key: str, aws_region: str = 'us-east-1'):
        """Initialize the parser with OpenAI API key and AWS credentials"""
        self.client = OpenAI(api_key=openai_api_key)
        
        # Initialize AWS Textract client
        self.textract_client = boto3.client(
            'textract',
            aws_access_key_id=aws_access_key,
            aws_secret_access_key=aws_secret_key,
            region_name=aws_region
        )
        
        # Enhanced prompt for medical document parsing with classification
        self.medical_parsing_prompt = """
        You are an expert medical document parser, Named Entity Recognition (NER) specialist, and document classifier. 
        Analyze the provided medical document/image and extract ALL relevant patient information.
        
        Please extract and return a JSON object with the following structure:
        
        {
            "document_classification": {
                "document_type": "Choose from: [lab_reports, prescription, insurance_medical_request_forms, discharge_summary]",
                "confidence_score": "Rate confidence 0-100% for classification",
                "reasoning": "Brief explanation why this classification was chosen",
                "secondary_type": "Alternative classification if uncertain"
            },
            "document_sentiment": {
                "overall_tone": "professional/urgent/routine/concerning/positive/negative",
                "urgency_level": "low/medium/high/critical",
                "completeness": "complete/incomplete/partial",
                "legibility": "clear/somewhat_clear/poor/illegible",
                "document_condition": "excellent/good/fair/poor/damaged"
            },
            "patient_info": {
                "name": "Full patient name",
                "age": "Patient age",
                "gender": "M/F/Other",
                "date_of_birth": "DOB if available",
                "contact_number": "Phone number",
                "email": "Email address",
                "address": "Patient address"
            },
            "medical_details": {
                "medical_record_number": "MRN or patient ID",
                "visit_date": "Date of visit/report",
                "doctor_name": "Attending physician",
                "department": "Department/specialty",
                "diagnosis": "Primary diagnosis",
                "secondary_diagnosis": "Additional diagnoses",
                "symptoms": ["List of symptoms"],
                "medications": ["List of medications"],
                "allergies": ["Known allergies"],
                "vital_signs": {
                    "blood_pressure": "BP reading",
                    "heart_rate": "HR",
                    "temperature": "Temperature",
                    "weight": "Weight",
                    "height": "Height"
                }
            },
            "laboratory_results": {
                "test_name": "Name of test/lab work",
                "test_date": "Date of test",
                "results": ["List of test results with values"],
                "reference_ranges": ["Normal ranges if provided"],
                "abnormal_findings": ["Any abnormal results"],
                "critical_values": ["Any critical or panic values"]
            },
            "document_info": {
                "hospital_name": "Hospital/clinic name",
                "report_date": "Date of report generation",
                "lab_number": "Lab/report number if applicable",
                "page_number": "Page number if multi-page",
                "document_id": "Any document identifier"
            },
            "additional_entities": {
                "procedures": ["Medical procedures mentioned"],
                "follow_up": "Follow-up instructions",
                "next_appointment": "Next appointment date",
                "insurance_info": "Insurance details if available",
                "emergency_contact": "Emergency contact info",
                "provider_notes": "Doctor's notes or comments",
                "billing_codes": ["ICD codes, CPT codes if present"]
            },
            "quality_assessment": {
                "information_completeness": "Rate 0-100% how complete the document is",
                "data_reliability": "high/medium/low based on document quality",
                "missing_critical_info": ["List any critical missing information"],
                "extraction_challenges": ["Any OCR or parsing difficulties encountered"]
            }
        }
        
        CLASSIFICATION GUIDELINES:
        - lab_reports: Blood tests, urine tests, imaging reports, pathology reports, diagnostic test results
        - prescription: Medication prescriptions, pharmacy orders, drug recommendations
        - insurance_medical_request_forms: Prior authorization forms, medical necessity forms, insurance claims
        - discharge_summary: Hospital discharge papers, treatment summaries, care instructions
        
        IMPORTANT INSTRUCTIONS:
        1. If any field is not found in the document, use null instead of empty string
        2. Extract dates in DD/MM/YYYY or MM/DD/YYYY format consistently  
        3. For numerical values, include units (e.g., "120/80 mmHg", "98.6°F")
        4. Be thorough - extract even partial information
        5. If text is unclear due to OCR, indicate with [UNCLEAR] but still provide best guess
        6. Return ONLY the JSON object, no additional text
        7. Ensure the JSON is properly formatted and valid
        8. Be very careful with document classification - look for key indicators
        9. Assess document quality and note any extraction challenges
        
        Analyze the document thoroughly and extract all available medical information with classification.
        """
    
    def detect_file_type(self, file_path: str) -> str:
        """Detect the type of uploaded file using multiple methods"""
        try:
            # Method 1: Use mimetypes library (built-in Python)
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type:
                if 'image' in mime_type:
                    return 'image'
                elif 'pdf' in mime_type:
                    return 'pdf'
                elif 'word' in mime_type or 'officedocument' in mime_type:
                    return 'word'
            
            # Method 2: Check file extension
            extension = Path(file_path).suffix.lower()
            if extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                return 'image'
            elif extension == '.pdf':
                return 'pdf'
            elif extension in ['.docx', '.doc']:
                return 'word'
            
            # Method 3: Try to read file headers (magic numbers)
            with open(file_path, 'rb') as f:
                header = f.read(16)
                
            # Check for common file signatures
            if header.startswith(b'\x89PNG'):
                return 'image'
            elif header.startswith(b'\xff\xd8\xff'):  # JPEG
                return 'image'
            elif header.startswith(b'%PDF'):
                return 'pdf'
            elif header.startswith(b'PK\x03\x04') and extension in ['.docx', '.doc']:
                return 'word'
            
            return 'unknown'
            
        except Exception as e:
            print(f"File type detection error: {e}")
            # Final fallback - just use extension
            extension = Path(file_path).suffix.lower()
            if extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']:
                return 'image'
            elif extension == '.pdf':
                return 'pdf'
            elif extension in ['.docx', '.doc']:
                return 'word'
            else:
                return 'unknown'
    
    def encode_image_to_base64(self, image_path: str) -> str:
        """Convert image to base64 for OpenAI API"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def extract_text_with_textract(self, file_path: str, file_type: str) -> str:
        """Extract text using AWS Textract"""
        try:
            with open(file_path, 'rb') as document:
                document_bytes = document.read()
            
            if file_type == 'image':
                # For images, use detect_document_text
                response = self.textract_client.detect_document_text(
                    Document={'Bytes': document_bytes}
                )
            elif file_type == 'pdf':
                # For PDFs, use detect_document_text (single page) or start_document_text_detection for multi-page
                response = self.textract_client.detect_document_text(
                    Document={'Bytes': document_bytes}
                )
            else:
                return ""
            
            # Extract text from Textract response
            extracted_text = ""
            for item in response['Blocks']:
                if item['BlockType'] == 'LINE':
                    extracted_text += item['Text'] + '\n'
            
            return extracted_text
            
        except ClientError as e:
            print(f"AWS Textract error: {e}")
            return ""
        except Exception as e:
            print(f"Textract extraction failed: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: str) -> tuple:
        """Extract text and convert first page to image from PDF"""
        text_content = ""
        first_page_image = None
        
        try:
            # First try Textract for PDF
            textract_text = self.extract_text_with_textract(pdf_path, 'pdf')
            if textract_text:
                text_content = textract_text
            else:
                # Fallback to PyPDF2
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n"
            
            # Convert first page to image for vision API
            images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=200)
            if images:
                first_page_image = images[0]
                # Save temporarily
                temp_image_path = "temp_pdf_page.png"
                first_page_image.save(temp_image_path, 'PNG')
                return text_content, temp_image_path
                
        except Exception as e:
            print(f"PDF processing error: {e}")
        
        return text_content, None
    
    def extract_text_from_word(self, word_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(word_path)
            text_content = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            return text_content
        except Exception as e:
            print(f"Word document processing error: {e}")
            return ""
    
    def parse_with_openai_vision(self, image_path: str, extracted_text: str = "") -> Dict[str, Any]:
        """Use OpenAI Vision API to parse medical document"""
        try:
            base64_image = self.encode_image_to_base64(image_path)
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": self.medical_parsing_prompt + 
                                   (f"\n\nTextract extracted text for reference:\n{extracted_text}" if extracted_text else "")
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}",
                                "detail": "high"
                            }
                        }
                    ]
                }
            ]
            
            response = self.client.chat.completions.create(
                model="gpt-4o",  # Using GPT-4 with vision
                messages=messages,
                max_tokens=3000,
                temperature=0.1
            )
            
            result_text = response.choices[0].message.content
            
            # Try to parse as JSON
            try:
                return json.loads(result_text)
            except json.JSONDecodeError:
                # If JSON parsing fails, try to extract JSON from the response
                import re
                json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
                else:
                    return {"error": "Could not parse JSON response", "raw_response": result_text}
                    
        except Exception as e:
            return {"error": f"OpenAI API error: {str(e)}"}
    
    def parse_with_openai_text(self, text_content: str) -> Dict[str, Any]:
        """Use OpenAI to parse text-only content"""
        try:
            messages = [
                {
                    "role": "user",
                    "content": self.medical_parsing_prompt + f"\n\nDocument text to analyze:\n{text_content}"
                }
            ]
            
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=messages,
                max_tokens=3000,
                temperature=0.1
            )
            
            result_text = response.choices[0].message.content
            
            # Try to parse as JSON
            try:
                return json.loads(result_text)
            except json.JSONDecodeError:
                # If JSON parsing fails, try to extract JSON from the response
                import re
                json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
                else:
                    return {"error": "Could not parse JSON response", "raw_response": result_text}
                    
        except Exception as e:
            return {"error": f"OpenAI API error: {str(e)}"}
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Main method to process any type of document"""
        print(f"Processing file: {file_path}")
        
        file_type = self.detect_file_type(file_path)
        print(f"Detected file type: {file_type}")
        
        result = {"file_type": file_type, "file_path": file_path}
        
        try:
            if file_type == 'image':
                # For images, use Textract + vision API
                textract_text = self.extract_text_with_textract(file_path, 'image')
                parsing_result = self.parse_with_openai_vision(file_path, textract_text)
                
            elif file_type == 'pdf':
                # For PDFs, extract text with Textract and convert first page to image
                text_content, image_path = self.extract_text_from_pdf(file_path)
                if image_path:
                    parsing_result = self.parse_with_openai_vision(image_path, text_content)
                    # Clean up temp file
                    try:
                        os.remove(image_path)
                    except:
                        pass
                else:
                    parsing_result = self.parse_with_openai_text(text_content)
                    
            elif file_type == 'word':
                # For Word docs, extract text only
                text_content = self.extract_text_from_word(file_path)
                parsing_result = self.parse_with_openai_text(text_content)
                
            else:
                return {"error": f"Unsupported file type: {file_type}"}
            
            # Merge results
            result.update(parsing_result)
            return result
            
        except Exception as e:
            return {"error": f"Processing error: {str(e)}"}

# Main execution function
def main():
    """Main function to run in Google Colab"""
    
    # Get API keys
    print("Please enter your API credentials:")
    openai_api_key = input("OpenAI API Key: ").strip()
    aws_access_key = input("AWS Access Key ID: ").strip()
    aws_secret_key = input("AWS Secret Access Key: ").strip()
    aws_region = input("AWS Region (default: us-east-1): ").strip() or "us-east-1"
    
    if not openai_api_key or not aws_access_key or not aws_secret_key:
        print("❌ All API credentials are required!")
        return
    
    # Initialize parser
    try:
        parser = MedicalDocumentParser(openai_api_key, aws_access_key, aws_secret_key, aws_region)
        print("✅ Parser initialized successfully!")
    except Exception as e:
        print(f"❌ Failed to initialize parser: {e}")
        return
    
    print("\n" + "="*70)
    print("🏥 ENHANCED MEDICAL DOCUMENT PARSER WITH AWS TEXTRACT")
    print("Supports: Images (JPG, PNG), PDFs, Word Documents")
    print("Features: Document Classification, Sentiment Analysis, OCR with Textract")
    print("="*70)
    
    while True:
        print("\n📁 Upload your medical document:")
        
        # Upload file
        uploaded = files.upload()
        
        if not uploaded:
            print("No file uploaded.")
            continue
        
        # Process each uploaded file
        for filename, content in uploaded.items():
            print(f"\n🔍 Processing: {filename}")
            
            # Save uploaded file temporarily
            with open(filename, 'wb') as f:
                f.write(content)
            
            # Process the document
            result = parser.process_document(filename)
            
            # Display results with better formatting
            print("\n" + "="*60)
            print("📋 EXTRACTION RESULTS:")
            print("="*60)
            
            # Display key information first
            if 'document_classification' in result:
                print("🏷️  DOCUMENT CLASSIFICATION:")
                classification = result['document_classification']
                print(f"   Type: {classification.get('document_type', 'Unknown')}")
                print(f"   Confidence: {classification.get('confidence_score', 'N/A')}")
                print(f"   Reasoning: {classification.get('reasoning', 'N/A')}")
                print()
            
            if 'document_sentiment' in result:
                print("📊 DOCUMENT ASSESSMENT:")
                sentiment = result['document_sentiment']
                print(f"   Overall Tone: {sentiment.get('overall_tone', 'N/A')}")
                print(f"   Urgency Level: {sentiment.get('urgency_level', 'N/A')}")
                print(f"   Completeness: {sentiment.get('completeness', 'N/A')}")
                print(f"   Legibility: {sentiment.get('legibility', 'N/A')}")
                print()
            
            # Full JSON output
            print("📄 COMPLETE RESULTS:")
            print(json.dumps(result, indent=2, ensure_ascii=False))
            
            # Save results to file
            output_filename = f"parsed_{filename}_results.json"
            with open(output_filename, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            
            print(f"\n💾 Results saved to: {output_filename}")
            
            # Download results
            files.download(output_filename)
            
            # Clean up
            try:
                os.remove(filename)
                os.remove(output_filename)
            except:
                pass
        
        # Ask if user wants to process another file
        another = input("\n🔄 Process another document? (y/n): ").strip().lower()
        if another != 'y':
            break
    
    print("\n✅ Processing complete!")

# Example usage function for testing
def test_with_sample():
    """Test function with sample data"""
    
    # Sample test - you need to provide your keys
    openai_key = "your-openai-api-key-here"
    aws_access = "your-aws-access-key-here"
    aws_secret = "your-aws-secret-key-here"
    
    parser = MedicalDocumentParser(openai_key, aws_access, aws_secret)
    
    # Test with uploaded file
    print("Upload a test file and we'll parse it:")
    uploaded = files.upload()
    
    for filename, content in uploaded.items():
        with open(filename, 'wb') as f:
            f.write(content)
        
        result = parser.process_document(filename)
        print("\nParsing Result:")
        print(json.dumps(result, indent=2))
        
        os.remove(filename)

# Run the main function
if __name__ == "__main__":
    main()

# Alternative: Uncomment below to run test function
# test_with_sample()